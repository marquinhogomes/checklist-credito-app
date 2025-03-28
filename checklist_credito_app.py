
import pandas as pd
import streamlit as st
from fpdf import FPDF
from docx import Document

csv_path = "Checklist_Analise_Credito_PowerBI.csv"

class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, "Relatório do Checklist de Crédito", ln=True, align="C")
        self.ln(10)

def gerar_pdf(df):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    for secao in df["Seção"].unique():
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, f"Seção: {secao}", ln=True)
        pdf.set_font("Arial", size=10)
        for _, row in df[df["Seção"] == secao].iterrows():
            pdf.multi_cell(0, 8, f"- {row['Item de Verificação']} | Status: {row['Status']} | Obs: {row['Observações']}")
        pdf.ln(5)
    pdf.output("Relatorio_Checklist_Credito.pdf")

def gerar_word(df):
    doc = Document()
    doc.add_heading("Relatório do Checklist de Crédito", 0)
    for secao in df["Seção"].unique():
        doc.add_heading(f"Seção: {secao}", level=1)
        for _, row in df[df["Seção"] == secao].iterrows():
            doc.add_paragraph(f"- {row['Item de Verificação']} | Status: {row['Status']} | Obs: {row['Observações']}")
    doc.save("Relatorio_Checklist_Credito.docx")

def main():
    st.title("Checklist de Análise de Crédito - Empresas")
    df = pd.read_csv(csv_path)
    secoes = df["Seção"].unique()
    secao_escolhida = st.selectbox("Escolha a seção para preencher:", secoes)
    df_secao = df[df["Seção"] == secao_escolhida].copy()
    for idx, row in df_secao.iterrows():
        status = st.selectbox(f"{row['Item de Verificação']}", ["", "OK", "Pendente", "N/A"], key=f"status_{idx}")
        obs = st.text_input(f"Observação (opcional)", value=row["Observações"], key=f"obs_{idx}")
        df.loc[idx, "Status"] = status
        df.loc[idx, "Observações"] = obs
    if st.button("Salvar checklist atualizado"):
        df.to_csv(csv_path, index=False)
        st.success("Checklist salvo com sucesso!")
    if st.checkbox("Mostrar checklist completo"):
        st.dataframe(df)
    st.markdown("---")
    st.subheader("Gerar Relatórios")
    if st.button("Gerar PDF e Word"):
        gerar_pdf(df)
        gerar_word(df)
        st.success("Relatórios PDF e Word gerados com sucesso!")

if __name__ == "__main__":
    main()
